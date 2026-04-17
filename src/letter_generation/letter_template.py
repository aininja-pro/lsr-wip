"""
Lien Letter Template Module

Generates pre-lien notice .docx letters using python-docx.
Each letter follows the approved LSR Multifamily template format.
"""

import io
import re
import logging
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

FONT_NAME = 'Calibri'
FONT_SIZE = Pt(11)


def _set_run_font(run, bold=False):
    """Apply standard font to a run."""
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE
    run.bold = bold


def _add_hyperlink(paragraph, url, text):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run_elem = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(FONT_SIZE.pt * 2)))
    rPr.append(sz)

    # Blue color
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    # Underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    run_elem.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


BODY_SPACING = Pt(8)  # Space after body paragraphs (matches Word default)
SECTION_BREAK = Pt(16)  # Larger gap between letter sections


def _add_paragraph(doc, text='', bold=False, space_after=None):
    """Add a paragraph with standard font formatting."""
    para = doc.add_paragraph()
    if text:
        run = para.add_run(text)
        _set_run_font(run, bold=bold)
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    return para


def _add_mixed_paragraph(doc, segments, space_after=None):
    """
    Add a paragraph with mixed bold/regular segments.

    segments: list of (text, bold) tuples
    """
    para = doc.add_paragraph()
    for text, bold in segments:
        run = para.add_run(text)
        _set_run_font(run, bold=bold)
    if space_after is not None:
        para.paragraph_format.space_after = space_after
    return para


def _format_date(d: date) -> str:
    """Format a date as 'February 16, 2026'."""
    return d.strftime('%B %d, %Y').replace(' 0', ' ')


def sanitize_filename(name: str) -> str:
    """
    Sanitize a string for use in a filename.
    Replace spaces with underscores, strip non-alphanumeric/underscore/hyphen chars.
    """
    name = name.replace(' ', '_')
    return re.sub(r'[^\w\-]', '', name)


def make_letter_filename(customer_name: str, invoice_number: str,
                         recipient_type: str = 'Customer',
                         extension: str = 'docx') -> str:
    """Generate the filename for a lien notice letter.

    Format: Lien_Notice_{customer_name}_{invoice_number}_{recipient_type}.{ext}
    """
    safe_name = sanitize_filename(customer_name or 'Unknown')
    safe_inv = sanitize_filename(invoice_number or 'NO_INV')
    safe_type = sanitize_filename(recipient_type or 'Customer') or 'Customer'
    return f"Lien_Notice_{safe_name}_{safe_inv}_{safe_type}.{extension}"


# Map recipient_type to the row-dict keys that describe their address block.
# Customer has no suite field; Manager and Owner do.
ADDRESS_FIELDS_BY_RECIPIENT = {
    'Customer': {
        'name': 'customer_name',
        'address': 'service_address',
        'suite': None,
        'city': 'service_city',
        'state': 'service_state',
        'zip': 'service_zip',
    },
    'Manager': {
        'name': 'manager_name',
        'address': 'manager_address',
        'suite': 'manager_suite',
        'city': 'manager_city',
        'state': 'manager_state',
        'zip': 'manager_zip',
    },
    'Owner': {
        'name': 'owner_name',
        'address': 'owner_address',
        'suite': 'owner_suite',
        'city': 'owner_city',
        'state': 'owner_state',
        'zip': 'owner_zip',
    },
}


def _is_blank(value) -> bool:
    """Return True for None, NaN, empty/whitespace-only strings."""
    if value is None:
        return True
    try:
        import math
        if isinstance(value, float) and math.isnan(value):
            return True
    except Exception:
        pass
    if isinstance(value, str) and value.strip() == '':
        return True
    return False


def format_address_line_with_suite(address: str, suite) -> str:
    """Combine address + suite as '{address}, {suite}' or just address when blank."""
    addr = (address or '').strip()
    if _is_blank(suite):
        return addr
    return f"{addr}, {str(suite).strip()}"


def generate_letter(row: dict, recipient_type: str, letter_date: date,
                    deadline_date: date, logo_path: str) -> io.BytesIO:
    """
    Generate a single pre-lien notice letter as a .docx file.

    Args:
        row: dict with canonical Phase 2 keys (customer_name, service_*,
             manager_*, owner_*, invoice_number, invoice_total, ...)
        recipient_type: "Customer", "Manager", or "Owner" — selects the
             address block fields. Everything else in the letter is identical
             across variants.
        letter_date: date to print on the letter
        deadline_date: payment deadline date
        logo_path: path to lsr_logo.png

    Returns:
        BytesIO containing the .docx file
    """
    if recipient_type not in ADDRESS_FIELDS_BY_RECIPIENT:
        raise ValueError(
            f"recipient_type must be one of "
            f"{sorted(ADDRESS_FIELDS_BY_RECIPIENT)}, got {recipient_type!r}"
        )
    doc = Document()

    # Page setup: US Letter, 1" margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default style: tight spacing (address blocks, signature need zero gaps)
    # Body paragraphs get explicit space_after where needed
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # RE: block always references the Customer + service location, regardless
    # of which recipient the letter is addressed to.
    customer_name = row.get('customer_name') or 'Unknown Property'
    service_address = row.get('service_address') or ''
    service_city = row.get('service_city') or ''
    service_state = row.get('service_state') or ''
    service_zip = row.get('service_zip') or ''
    invoice_number = row.get('invoice_number') or 'NO_INV'
    invoice_total = float(row.get('invoice_total') or 0)

    service_city_state_zip = f"{service_city}, {service_state} {service_zip}".strip()
    full_location = f"{service_address}, {service_city_state_zip}"

    # Recipient-specific address block fields.
    fields = ADDRESS_FIELDS_BY_RECIPIENT[recipient_type]
    recipient_name = row.get(fields['name']) or ''
    recipient_address = row.get(fields['address']) or ''
    recipient_suite = row.get(fields['suite']) if fields['suite'] else None
    recipient_city = row.get(fields['city']) or ''
    recipient_state = row.get(fields['state']) or ''
    recipient_zip = row.get(fields['zip']) or ''

    recipient_address_line = format_address_line_with_suite(
        recipient_address, recipient_suite
    )
    recipient_city_state_zip = (
        f"{recipient_city}, {recipient_state} {recipient_zip}".strip()
    )

    # 1. Logo
    logo = Path(logo_path)
    if logo.exists():
        para = doc.add_paragraph()
        para.add_run().add_picture(str(logo), width=Inches(4.5))
        para.paragraph_format.space_after = SECTION_BREAK
    else:
        logger.warning(f"Logo not found at {logo_path}")

    # 2. Date (bold)
    _add_paragraph(doc, _format_date(letter_date), bold=True, space_after=SECTION_BREAK)

    # 3. Recipient address block (tight — no spacing between lines).
    # Suite, when present, sits on the same line as the address, comma-separated.
    _add_paragraph(doc, recipient_name)
    _add_paragraph(doc, recipient_address_line)
    _add_paragraph(doc, recipient_city_state_zip, space_after=SECTION_BREAK)

    # 4. RE: block (always references the customer/service location)
    _add_paragraph(doc, 'RE:', bold=True)
    _add_paragraph(doc, f'Property: {customer_name}')
    _add_paragraph(doc, f'Location: {full_location}', space_after=SECTION_BREAK)

    # 5. Salutation
    _add_paragraph(doc, 'To Whom It May Concern:', bold=True, space_after=BODY_SPACING)

    # 6. Body paragraph 1
    _add_paragraph(
        doc,
        'We are writing to formally notify you that our company has provided '
        'labor and materials for the above-referenced property. According to '
        'our records, the following invoice remains unpaid:',
        space_after=BODY_SPACING,
    )

    # 7. Invoice bullet
    bullet_text = f"Invoice #{invoice_number} - ${invoice_total:,.2f}"
    bullet_para = doc.add_paragraph(style='List Bullet')
    run = bullet_para.add_run(bullet_text)
    _set_run_font(run, bold=True)
    bullet_para.paragraph_format.space_after = BODY_SPACING

    # 8. Body paragraph 2 (deadline date bold inline)
    deadline_str = _format_date(deadline_date)
    _add_mixed_paragraph(doc, [
        ("This notice is provided in accordance with the Texas mechanic\u2019s "
         "lien laws to protect our legal rights. Payment in full must be "
         "received in our office no later than ", False),
        (deadline_str, True),
        (", to avoid further action.", False),
    ], space_after=BODY_SPACING)

    # 9. Body paragraph 3 ($250.00 bold inline)
    _add_mixed_paragraph(doc, [
        ("If payment is not received by the deadline stated, we will proceed "
         "with filing a lien against the property, and an additional ", False),
        ("$250.00", True),
        (" in collection costs will be added to the balance due.", False),
    ], space_after=BODY_SPACING)

    # 10. Body paragraph 4
    _add_paragraph(
        doc,
        'We remain committed to resolving this matter amicably and are '
        'available to provide any additional documentation or information '
        'you may require. Please contact our office at your earliest '
        'convenience to confirm payment arrangements.',
        space_after=SECTION_BREAK,
    )

    # 11. Signature block (tight within)
    _add_paragraph(doc, 'Sincerely,', space_after=SECTION_BREAK)
    _add_paragraph(doc, 'Stephanie Campbell', bold=True)
    _add_paragraph(doc, 'AR Specialist')
    _add_paragraph(doc, 'LSR Multifamily')
    email_para = _add_paragraph(doc)
    _add_hyperlink(email_para, 'mailto:sstorey@lsrusa.com', 'sstorey@lsrusa.com')
    _add_paragraph(doc, '972-869-4479')

    # Save to BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
