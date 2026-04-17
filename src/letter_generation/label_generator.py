"""
Avery 5160 Mailing Label Generator

Produces a .docx file laid out as an Avery 5160 mailing label sheet:
- US Letter page (8.5" x 11")
- 30 labels per page (3 columns x 10 rows)
- Label size: 2.625" wide x 1" tall
- Top/bottom margins: 0.5"
- Side margins: 0.1875"
- Horizontal gap between labels: 0.125" (handled with a narrow spacer column)
- No vertical gap (labels are contiguous)

For each invoice row, labels are emitted in this order:
  1. 3 copies of Customer label
  2. 3 copies of Manager label
  3. 3 copies of Owner label (only when row['owner_letter_needed'] is True)

Labels fill left-to-right, top-to-bottom. A new page starts when the previous
page is full (30 cells).
"""

import io
import logging
import math
from typing import List, Dict

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .letter_template import (
    ADDRESS_FIELDS_BY_RECIPIENT,
    format_address_line_with_suite,
    _is_blank,
)

logger = logging.getLogger(__name__)

LABEL_FONT_NAME = 'Calibri'
LABEL_FONT_SIZE = Pt(10)

# Avery 5160 dimensions (inches).
LABEL_WIDTH = Inches(2.625)
LABEL_HEIGHT = Inches(1.0)
LABEL_HORIZONTAL_GAP = Inches(0.125)
PAGE_TOP_MARGIN = Inches(0.5)
PAGE_BOTTOM_MARGIN = Inches(0.5)
PAGE_SIDE_MARGIN = Inches(0.1875)

LABELS_PER_ROW = 3
ROWS_PER_PAGE = 10
LABELS_PER_PAGE = LABELS_PER_ROW * ROWS_PER_PAGE  # 30


def _build_label_text(row: Dict, recipient_type: str) -> List[str]:
    """Build the 2-3 text lines for a single label cell."""
    fields = ADDRESS_FIELDS_BY_RECIPIENT[recipient_type]
    name = row.get(fields['name']) or ''
    address = row.get(fields['address']) or ''
    suite = row.get(fields['suite']) if fields['suite'] else None
    city = row.get(fields['city']) or ''
    state = row.get(fields['state']) or ''
    zip_code = row.get(fields['zip']) or ''

    address_line = format_address_line_with_suite(address, suite)
    city_state_zip = f"{city}, {state} {zip_code}".strip()

    lines = [str(name).strip()] if not _is_blank(name) else []
    if not _is_blank(address_line):
        lines.append(address_line)
    if not _is_blank(city_state_zip) and city_state_zip != ',':
        lines.append(city_state_zip)
    return lines


def _collect_labels(rows: pd.DataFrame) -> List[List[str]]:
    """Walk rows and produce the ordered list of label blocks (each a list of text lines)."""
    labels: List[List[str]] = []
    for _, row in rows.iterrows():
        row_dict = row.to_dict()
        # 3 copies each of Customer, then Manager, then Owner (if needed).
        recipients = ['Customer', 'Manager']
        if row_dict.get('owner_letter_needed'):
            recipients.append('Owner')
        for recipient in recipients:
            text_lines = _build_label_text(row_dict, recipient)
            for _ in range(3):
                labels.append(list(text_lines))
    return labels


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    """Set cell padding in twips (1/20 pt). Small values keep content inside label."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('left', left), ('bottom', bottom), ('right', right)):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def _remove_all_borders(table):
    """Strip visible borders from a table so it prints as labels, not a grid."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tbl_pr.append(borders)


def _set_row_exact_height(row, height):
    """Pin a table row to an exact height (prevents auto-resize)."""
    tr_pr = row._tr.get_or_add_trPr()
    tr_height = OxmlElement('w:trHeight')
    tr_height.set(qn('w:val'), str(int(height)))
    tr_height.set(qn('w:hRule'), 'exact')
    tr_pr.append(tr_height)


def _fill_label_cell(cell, text_lines: List[str]):
    """Write the address lines into a label cell with Avery-appropriate formatting."""
    cell.text = ''  # clear default empty paragraph
    # Re-use the first paragraph docx creates.
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text_lines):
        if i == 0:
            run = para.add_run(line)
        else:
            new_para = cell.add_paragraph()
            new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            new_para.paragraph_format.space_before = Pt(0)
            new_para.paragraph_format.space_after = Pt(0)
            run = new_para.add_run(line)
        run.font.name = LABEL_FONT_NAME
        run.font.size = LABEL_FONT_SIZE


def _remove_default_first_paragraph(doc: Document):
    """python-docx creates an empty paragraph at the top of every document.
    Even with 0pt spacing, that paragraph's line height still takes ~0.17"
    — enough to push the 10th label row off the first page and cascade."""
    if not doc.paragraphs:
        return
    p = doc.paragraphs[0]
    parent = p._element.getparent()
    if parent is not None:
        parent.remove(p._element)


def generate_labels(rows: pd.DataFrame) -> io.BytesIO:
    """
    Generate an Avery 5160 label sheet as a .docx file.

    Args:
        rows: parsed invoice DataFrame (must include `owner_letter_needed`).

    Returns:
        BytesIO containing the .docx.
    """
    labels = _collect_labels(rows)
    n_labels = len(labels)
    n_pages = max(1, math.ceil(n_labels / LABELS_PER_PAGE)) if n_labels else 1
    logger.info(f"Generating {n_labels} labels across {n_pages} pages")

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = PAGE_TOP_MARGIN
    section.bottom_margin = PAGE_BOTTOM_MARGIN
    section.left_margin = PAGE_SIDE_MARGIN
    section.right_margin = PAGE_SIDE_MARGIN
    # Word reserves header/footer zones inside the page by default (0.5" each).
    # On an Avery 5160 sheet those zones cannibalize the last label row. Zero
    # them so the body uses the full area between top and bottom margins.
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    # Default style: tight spacing.
    style = doc.styles['Normal']
    style.font.name = LABEL_FONT_NAME
    style.font.size = LABEL_FONT_SIZE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    if not labels:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    # One continuous table spanning all pages. Word flows rows across pages
    # automatically, so we don't risk a leading paragraph's line height
    # displacing rows and pushing row 10 onto the next page.
    total_rows = math.ceil(n_labels / LABELS_PER_ROW)
    table = doc.add_table(rows=total_rows, cols=5)
    table.autofit = False
    _remove_all_borders(table)

    # Column widths: label, gap, label, gap, label.
    col_widths = [LABEL_WIDTH, LABEL_HORIZONTAL_GAP, LABEL_WIDTH,
                  LABEL_HORIZONTAL_GAP, LABEL_WIDTH]
    for col_idx, width in enumerate(col_widths):
        for tr in table.rows:
            tr.cells[col_idx].width = width

    # Pin every label row to exactly 1" (1440 twips). Exact heights prevent
    # Word from reflowing and keep the grid aligned with the Avery sheet.
    for tr in table.rows:
        _set_row_exact_height(tr, 1440)

    # Populate left-to-right, top-to-bottom. Do NOT set cell margins:
    # LibreOffice adds them on top of hRule="exact" row height, which pushes
    # the 10th row off the first page and cascades to an extra page per sheet.
    label_cols = [0, 2, 4]
    label_cursor = 0
    for r in range(total_rows):
        for c in label_cols:
            if label_cursor >= n_labels:
                break
            _fill_label_cell(table.rows[r].cells[c], labels[label_cursor])
            label_cursor += 1

    # Strip the default leading paragraph so the table starts flush with
    # the top margin — otherwise the first page fits only 9 rows.
    _remove_default_first_paragraph(doc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def count_labels(rows: pd.DataFrame) -> int:
    """Total labels a given DataFrame will produce (useful for UI summary)."""
    return len(_collect_labels(rows))


def count_label_pages(rows: pd.DataFrame) -> int:
    """Pages the label document will span (useful for UI summary)."""
    n = count_labels(rows)
    return max(1, math.ceil(n / LABELS_PER_PAGE)) if n else 0
